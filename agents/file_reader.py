"""文件读取工具 - 支持 PDF、DOCX、TXT 等格式

提供底层实现（_read_pdf、_read_docx 等），供 tools.py 注册为 agent-as-tool。
"""

from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# 支持的文件扩展名
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}


def _read_pdf(file_path: str) -> dict[str, Any]:
    """读取 PDF 文件，提取文本内容

    Args:
        file_path: PDF 文件路径

    Returns:
        {"filename": str, "content": str, "page_count": int, "format": "pdf"}
    """
    try:
        import fitz  # pymupdf
    except ImportError:
        return {
            "filename": os.path.basename(file_path),
            "content": "",
            "page_count": 0,
            "format": "pdf",
            "error": "pymupdf 未安装，请运行 pip install pymupdf",
        }

    doc = fitz.open(file_path)
    pages = []
    for page_num, page in enumerate(doc, 1):
        text = page.get_text()
        if text.strip():
            pages.append(f"--- 第 {page_num} 页 ---\n{text}")

    content = "\n\n".join(pages)
    result = {
        "filename": os.path.basename(file_path),
        "content": content,
        "page_count": len(doc),
        "format": "pdf",
    }
    doc.close()
    return result


def _read_docx(file_path: str) -> dict[str, Any]:
    """读取 DOCX 文件，提取文本内容

    Args:
        file_path: DOCX 文件路径

    Returns:
        {"filename": str, "content": str, "paragraph_count": int, "format": "docx"}
    """
    try:
        from docx import Document
    except ImportError:
        return {
            "filename": os.path.basename(file_path),
            "content": "",
            "paragraph_count": 0,
            "format": "docx",
            "error": "python-docx 未安装，请运行 pip install python-docx",
        }

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(paragraphs)

    # 也提取表格内容
    tables_text = []
    for i, table in enumerate(doc.tables, 1):
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        if rows_text:
            tables_text.append(f"--- 表格 {i} ---\n" + "\n".join(rows_text))

    if tables_text:
        content += "\n\n" + "\n\n".join(tables_text)

    return {
        "filename": os.path.basename(file_path),
        "content": content,
        "paragraph_count": len(paragraphs),
        "format": "docx",
    }


def _read_txt(file_path: str) -> dict[str, Any]:
    """读取纯文本文件（TXT / MD 等）

    Args:
        file_path: 文本文件路径

    Returns:
        {"filename": str, "content": str, "line_count": int, "format": "txt"}
    """
    # 尝试多种编码
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()
            break
        except UnicodeDecodeError:
            continue
    else:
        content = ""

    return {
        "filename": os.path.basename(file_path),
        "content": content,
        "line_count": content.count("\n") + 1 if content else 0,
        "format": "txt",
    }


def read_file(file_path: str) -> dict[str, Any]:
    """根据文件扩展名自动选择读取方式

    Args:
        file_path: 文件路径

    Returns:
        统一结构的文件内容字典：
        {"filename": str, "content": str, "format": str, ...}
        如果读取失败，包含 "error" 字段
    """
    path = Path(file_path)
    if not path.exists():
        return {
            "filename": os.path.basename(file_path),
            "content": "",
            "format": "unknown",
            "error": f"文件不存在: {file_path}",
        }

    ext = path.suffix.lower()

    if ext == ".pdf":
        return _read_pdf(file_path)
    elif ext in (".docx", ".doc"):
        if ext == ".doc":
            return {
                "filename": os.path.basename(file_path),
                "content": "",
                "format": "doc",
                "error": ".doc 格式不支持，请转换为 .docx",
            }
        return _read_docx(file_path)
    elif ext in (".txt", ".md"):
        return _read_txt(file_path)
    else:
        return {
            "filename": os.path.basename(file_path),
            "content": "",
            "format": ext,
            "error": f"不支持的文件格式: {ext}，支持的格式: {SUPPORTED_EXTENSIONS}",
        }


def read_files(file_paths: list[str]) -> list[dict[str, Any]]:
    """批量读取多个文件

    Args:
        file_paths: 文件路径列表

    Returns:
        文件内容字典列表
    """
    results = []
    for fp in file_paths:
        logger.info(f"[file_reader] 读取文件: {fp}")
        result = read_file(fp)
        if "error" in result:
            logger.warning(f"[file_reader] 读取失败: {result['error']}")
        else:
            logger.info(
                f"[file_reader] 读取成功: {result['filename']} "
                f"({result.get('page_count') or result.get('paragraph_count') or result.get('line_count', '?')} 单元)"
            )
        results.append(result)
    return results
